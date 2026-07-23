#!/usr/bin/env python3
"""
Generic, data-driven resume generator (DOCX).

Usage:
    python3 generate_resume.py --data <data.json> [--config <config.yaml>] [--out-dir <dir>]

The tailoring workflow writes a per-run `data.json` (company/role + the selected
experience/project/skill content) and points this script at the user's style
`config.yaml`. Nothing in this file is personal or needs hand-editing — all content
comes from --data and all style/rules from --config (with built-in fallbacks).
"""

import argparse
import json
import os
import re
import sys

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import docx.opc.constants


# =============================================================================
# CONFIG (built-in defaults; overridden by --config)
# =============================================================================

DEFAULT_CONFIG = {
    "font": "Garamond",
    "size": 9.5,
    "name_size": 14,
    "margins": 0.5,                 # inches, all sides
    "link_color": "0563C1",
    "word_target": {"min": 665, "max": 700},
    "fixed_word_overhead": 120,     # header, edu, role lines, proj names, skills
    "bullet_char_min": 160,
    "bullet_char_max": 220,
    "filename_pattern": "{name}_Resume_{company}{team}",
    "output_dir": "~/Documents/Resumes/{company}",
    "guardrails": {
        "seniority_cap": True,
        "em_dash_check": True,
        "formula_variety": True,
        "domain_gap_flag": True,
    },
}


def _deep_merge(base, override):
    out = dict(base)
    for k, v in (override or {}).items():
        if isinstance(v, dict) and isinstance(out.get(k), dict):
            out[k] = _deep_merge(out[k], v)
        else:
            out[k] = v
    return out


def load_config(path):
    """Load config from YAML or JSON, merged over DEFAULT_CONFIG. Missing file/dep -> defaults."""
    if not path:
        return dict(DEFAULT_CONFIG)
    path = os.path.expanduser(path)
    if not os.path.exists(path):
        print(f"[config] {path} not found; using built-in defaults.")
        return dict(DEFAULT_CONFIG)
    with open(path, "r") as f:
        raw = f.read()
    loaded = None
    if path.endswith(".json"):
        loaded = json.loads(raw)
    else:
        try:
            import yaml  # lazy: only needed for YAML configs
            loaded = yaml.safe_load(raw)
        except ImportError:
            print("[config] pyyaml not installed; using built-in defaults. "
                  "Install with: pip install pyyaml")
            return dict(DEFAULT_CONFIG)
    return _deep_merge(DEFAULT_CONFIG, loaded or {})


def _slug_name(name):
    parts = [p for p in re.split(r"\s+", name.strip()) if p]
    return "_".join(p[:1].upper() + p[1:] for p in parts) if parts else "Resume"


def _safe_component(text):
    return re.sub(r"[^\w.-]+", "_", text.strip()).strip("_") or "Unknown"


# =============================================================================
# --- FORMATTING ENGINE --- (rarely needs editing; driven by CFG + data)
# =============================================================================

def build_resume(data, cfg):
    FONT = cfg["font"]
    SZ = float(cfg["size"])
    NAME_SZ = float(cfg["name_size"])
    LINK_COLOR = cfg["link_color"]
    MARGIN = float(cfg["margins"])
    TEXT_WIDTH_TWIPS = int(7.5 * 1440)

    doc = Document()
    doc.styles["Normal"].paragraph_format.space_before = Pt(0)
    doc.styles["Normal"].paragraph_format.space_after = Pt(0)
    doc.styles["Normal"].font.name = FONT
    doc.styles["Normal"].font.size = Pt(SZ)

    sec = doc.sections[0]
    sec.left_margin = Inches(MARGIN)
    sec.right_margin = Inches(MARGIN)
    sec.top_margin = Inches(MARGIN)
    sec.bottom_margin = Inches(MARGIN)

    def _apply_font(rPr, bold=False, size=SZ):
        for tag in ["w:rFonts", "w:sz", "w:szCs"]:
            for el in rPr.findall(qn(tag)):
                rPr.remove(el)
        rf = OxmlElement("w:rFonts")
        rf.set(qn("w:ascii"), FONT)
        rf.set(qn("w:hAnsi"), FONT)
        rf.set(qn("w:cs"), FONT)
        rPr.insert(0, rf)
        sz_el = OxmlElement("w:sz")
        sz_el.set(qn("w:val"), str(int(size * 2)))
        szCs_el = OxmlElement("w:szCs")
        szCs_el.set(qn("w:val"), str(int(size * 2)))
        rPr.append(sz_el)
        rPr.append(szCs_el)
        if bold:
            for el in rPr.findall(qn("w:b")):
                rPr.remove(el)
            rPr.insert(0, OxmlElement("w:b"))

    def _para(sb=0, sa=2, align=WD_ALIGN_PARAGRAPH.LEFT, rtab=False):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(sb)
        p.paragraph_format.space_after = Pt(sa)
        p.paragraph_format.alignment = align
        if rtab:
            pPr = p._p.get_or_add_pPr()
            tabs = OxmlElement("w:tabs")
            tab = OxmlElement("w:tab")
            tab.set(qn("w:val"), "right")
            tab.set(qn("w:pos"), str(TEXT_WIDTH_TWIPS))
            tabs.append(tab)
            pPr.append(tabs)
        return p

    def _r(p, text, bold=False, size=SZ):
        run = p.add_run(text)
        run.bold = bold
        run.italic = False
        run.font.name = FONT
        run.font.size = Pt(size)
        _apply_font(run._r.get_or_add_rPr(), bold=bold, size=size)
        return run

    def _link(p, text, url, bold=False):
        rid = p.part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
        hl = OxmlElement("w:hyperlink")
        hl.set(qn("r:id"), rid)
        run_el = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        _apply_font(rPr, bold=bold)
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rPr.append(u)
        col = OxmlElement("w:color")
        col.set(qn("w:val"), LINK_COLOR)
        rPr.append(col)
        run_el.append(rPr)
        t = OxmlElement("w:t")
        t.text = text
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        run_el.append(t)
        hl.append(run_el)
        p._p.append(hl)

    def _section(title):
        p = _para(sb=6, sa=1)
        _r(p, title, bold=True)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bot = OxmlElement("w:bottom")
        bot.set(qn("w:val"), "single")
        bot.set(qn("w:sz"), "4")
        bot.set(qn("w:space"), "1")
        bot.set(qn("w:color"), "000000")
        pBdr.append(bot)
        pPr.append(pBdr)

    bullet_texts = []

    def _bul(text):
        bullet_texts.append(text)
        p = _para(sb=0, sa=1)
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.first_line_indent = Inches(-0.2)
        pPr = p._p.get_or_add_pPr()
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "left")
        tab.set(qn("w:pos"), "288")
        tabs.append(tab)
        pPr.append(tabs)
        _r(p, "•\t")
        _r(p, text)

    header = data["header"]
    education = data.get("education", [])
    experience = data.get("experience", [])
    projects = data.get("projects", [])
    skills = data.get("skills", [])

    # ── Header ───────────────────────────────────────────────────────
    p0 = _para(sb=0, sa=1, align=WD_ALIGN_PARAGRAPH.CENTER)
    run = p0.add_run(header["name"])
    _apply_font(run._r.get_or_add_rPr(), bold=True, size=NAME_SZ)
    run.bold = True
    run.italic = False
    run.font.size = Pt(NAME_SZ)
    run.font.name = FONT

    # Contact line: phone | email | any present links (linkedin/github/website/...)
    p0 = _para(sb=0, sa=(1 if header.get("auth") else 4), align=WD_ALIGN_PARAGRAPH.CENTER)
    contact_bits = [b for b in [header.get("phone"), header.get("email")] if b]
    _r(p0, " | ".join(contact_bits))
    link_keys = [k for k in ("linkedin", "github", "website") if header.get(k)]
    # include any additional link dicts not in the standard three
    for k in header:
        if k not in ("name", "phone", "email", "auth") and k not in link_keys \
                and isinstance(header.get(k), dict) and header[k].get("url"):
            link_keys.append(k)
    for link in link_keys:
        _r(p0, " | ")
        _link(p0, header[link]["text"], header[link]["url"])

    if header.get("auth"):
        p0 = _para(sb=0, sa=4, align=WD_ALIGN_PARAGRAPH.CENTER)
        _r(p0, header["auth"])

    # ── Education ────────────────────────────────────────────────────
    if education:
        _section("EDUCATION")
        for i, edu in enumerate(education):
            sb = 2 if i == 0 else 1
            detail_str = f' | {edu["detail"]}' if edu.get("detail") else ""
            p0 = _para(sb=sb, sa=0, rtab=True)
            _r(p0, edu["school"], bold=True)
            _r(p0, f' | {edu["degree"]}{detail_str}')
            _r(p0, "\t")
            _r(p0, edu.get("dates", ""))
            if edu.get("coursework"):
                p0 = _para(sb=0, sa=1)
                _r(p0, f'Relevant Coursework: {edu["coursework"]}')

    # ── Professional Experience ───────────────────────────────────────
    if experience:
        _section("PROFESSIONAL EXPERIENCE")
        for i, exp in enumerate(experience):
            sb = 2 if i == 0 else 3
            p0 = _para(sb=sb, sa=0, rtab=True)
            head_bits = [exp.get("company"), exp.get("role"), exp.get("location")]
            _r(p0, " | ".join(b for b in head_bits if b), bold=True)
            _r(p0, "\t")
            _r(p0, exp.get("dates", ""))
            for bullet in exp.get("bullets", []):
                _bul(bullet)

    # ── Projects ─────────────────────────────────────────────────────
    if projects:
        _section("PROJECTS")
        for proj in projects:
            p0 = _para(sb=2, sa=0)
            if proj.get("name_url"):
                _link(p0, proj["name"], proj["name_url"], bold=True)
            else:
                _r(p0, proj["name"], bold=True)
            if proj.get("stack"):
                _r(p0, f' | {proj["stack"]}')
            if proj.get("repo_url") and proj.get("repo_display"):
                _r(p0, " | ")
                _link(p0, proj["repo_display"], proj["repo_url"])
            for bullet in proj.get("bullets", []):
                _bul(bullet)

    # ── Technical Skills ─────────────────────────────────────────────
    if skills:
        _section("TECHNICAL SKILLS")
        for i, row in enumerate(skills):
            label, vals = row[0], row[1]
            p0 = _para(sb=(2 if i == 0 else 1), sa=0)
            _r(p0, f"{label}: ", bold=True)
            _r(p0, vals)

    return doc, bullet_texts


# =============================================================================
# Validation, filename, save
# =============================================================================

def report_and_save(doc, bullet_texts, data, cfg, out_dir_override=None):
    experience = data.get("experience", [])
    projects = data.get("projects", [])

    # Word-count estimate
    bullet_wc = sum(len(b.split()) for e in experience for b in e.get("bullets", [])) \
        + sum(len(b.split()) for p in projects for b in p.get("bullets", []))
    overhead = int(cfg["fixed_word_overhead"])
    est_total = bullet_wc + overhead
    wt = cfg["word_target"]
    print(f"Est. word count: {est_total} total  ({bullet_wc} bullet words + ~{overhead} fixed)"
          f"  |  target {wt['min']}-{wt['max']}")
    if est_total < wt["min"]:
        print(f"  -> BELOW target by ~{wt['min'] - est_total} words: consider a backfill bullet/project.")
    elif est_total > wt["max"]:
        print(f"  -> ABOVE target by ~{est_total - wt['max']} words: consider trimming.")

    # Bullet length validation
    cmin, cmax = int(cfg["bullet_char_min"]), int(cfg["bullet_char_max"])
    print(f"\nBullet validation (non-space char count, target {cmin}-{cmax}):")
    all_pass = True
    for i, text in enumerate(bullet_texts):
        count = len(text.replace(" ", ""))
        if count < cmin:
            status, ok = "SHORT", False
        elif count > cmax:
            status, ok = "LONG ", True   # long is a soft warning, not a hard fail
        else:
            status, ok = "OK   ", True
        if not ok:
            all_pass = False
        print(f"  [{status}] B{i+1:02d} ({count:3d}): {text[:60]}...")

    # Em-dash guardrail (mechanical check)
    if cfg.get("guardrails", {}).get("em_dash_check", True):
        offenders = [i + 1 for i, t in enumerate(bullet_texts) if "—" in t]
        if offenders:
            all_pass = False
            print(f"\n[EM-DASH] found em dashes in bullets: {offenders} — replace with comma/colon/period.")

    print("\nAll bullets passed." if all_pass
          else "\nWARNING: some bullets need attention before using (see above).")

    # Filename + output dir
    company = data.get("company", "Company")
    team = data.get("team", "")
    company_slug = _safe_component(company)
    team_suffix = f"_{_safe_component(team)}" if team else ""
    name_slug = _slug_name(data["header"]["name"])
    fname = cfg["filename_pattern"].format(name=name_slug, company=company_slug, team=team_suffix)
    if not fname.endswith(".docx"):
        fname += ".docx"

    out_dir = out_dir_override or cfg["output_dir"].format(company=company_slug)
    out_dir = os.path.expanduser(out_dir)
    os.makedirs(out_dir, exist_ok=True)
    out = os.path.join(out_dir, fname)
    doc.save(out)
    print(f"\nSaved: {out}")
    return out


def main():
    ap = argparse.ArgumentParser(description="Generate a tailored resume DOCX from a data file.")
    ap.add_argument("--data", required=True, help="Path to the resume data JSON.")
    ap.add_argument("--config", default=None, help="Path to a style/rules config (YAML or JSON).")
    ap.add_argument("--out-dir", default=None, help="Override output directory.")
    args = ap.parse_args()

    with open(os.path.expanduser(args.data), "r") as f:
        data = json.load(f)
    cfg = load_config(args.config)

    if "header" not in data or "name" not in data.get("header", {}):
        sys.exit("data file must contain header.name")

    doc, bullet_texts = build_resume(data, cfg)
    report_and_save(doc, bullet_texts, data, cfg, out_dir_override=args.out_dir)


if __name__ == "__main__":
    main()
