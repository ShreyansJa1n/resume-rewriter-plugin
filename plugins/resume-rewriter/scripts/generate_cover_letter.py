#!/usr/bin/env python3
"""
Generic, data-driven cover letter generator (DOCX).

Usage:
    python3 generate_cover_letter.py --data <data.json> [--config <config.yaml>] [--out-dir <dir>]

--data JSON shape:
    {
      "company": "Acme",
      "date": "July 17, 2026",
      "greeting": "Dear Acme Hiring Team,",
      "body": ["paragraph one", "paragraph two", ...],
      "closing": "Sincerely,",
      "signer": {
        "name": "Jane Doe", "phone": "...", "email": "...",
        "linkedin": {"text": "...", "url": "..."},   # optional
        "github":   {"text": "...", "url": "..."}    # optional
      }
    }

Style comes from --config (shares font/margins/link_color with the resume; uses its own
`cover_letter_size` and `cover_letter_filename_pattern`). Built-in fallbacks apply if absent.
"""

import argparse
import json
import os
import re
import sys

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import docx.opc.constants


DEFAULT_CONFIG = {
    "font": "Garamond",
    "cover_letter_size": 10,
    "name_size": 14,
    "margins": 0.5,
    "link_color": "0563C1",
    "cover_letter_filename_pattern": "{name}_CoverLetter_{company}",
    "output_dir": "~/Documents/Resumes/{company}",
}


def _deep_merge(base, override):
    out = dict(base)
    for k, v in (override or {}).items():
        if isinstance(v, dict) and isinstance(out.get(k), dict):
            out[k] = _deep_merge(out[k], v)
        else:
            out[k] = v
    return out


def load_config(path):
    if not path:
        return dict(DEFAULT_CONFIG)
    path = os.path.expanduser(path)
    if not os.path.exists(path):
        print(f"[config] {path} not found; using built-in defaults.")
        return dict(DEFAULT_CONFIG)
    with open(path, "r") as f:
        raw = f.read()
    if path.endswith(".json"):
        loaded = json.loads(raw)
    else:
        try:
            import yaml
            loaded = yaml.safe_load(raw)
        except ImportError:
            print("[config] pyyaml not installed; using built-in defaults.")
            return dict(DEFAULT_CONFIG)
    return _deep_merge(DEFAULT_CONFIG, loaded or {})


def _slug_name(name):
    parts = [p for p in re.split(r"\s+", name.strip()) if p]
    return "_".join(p[:1].upper() + p[1:] for p in parts) if parts else "Applicant"


def _safe_component(text):
    return re.sub(r"[^\w.-]+", "_", text.strip()).strip("_") or "Unknown"


def build_cover_letter(data, cfg):
    FONT = cfg["font"]
    SZ = float(cfg["cover_letter_size"])
    NAME_SZ = float(cfg["name_size"])
    LINK_COLOR = cfg["link_color"]
    MARGIN = float(cfg["margins"])

    signer = data["signer"]
    body = data.get("body", [])

    doc = Document()
    doc.styles["Normal"].paragraph_format.space_before = Pt(0)
    doc.styles["Normal"].paragraph_format.space_after = Pt(0)
    doc.styles["Normal"].font.name = FONT
    doc.styles["Normal"].font.size = Pt(SZ)

    sec = doc.sections[0]
    sec.left_margin = Inches(MARGIN)
    sec.right_margin = Inches(MARGIN)
    sec.top_margin = Inches(MARGIN)
    sec.bottom_margin = Inches(MARGIN)

    def _apply_font(rPr, bold=False, size=SZ):
        for tag in ["w:rFonts", "w:sz", "w:szCs"]:
            for el in rPr.findall(qn(tag)):
                rPr.remove(el)
        rf = OxmlElement("w:rFonts")
        rf.set(qn("w:ascii"), FONT)
        rf.set(qn("w:hAnsi"), FONT)
        rf.set(qn("w:cs"), FONT)
        rPr.insert(0, rf)
        sz_el = OxmlElement("w:sz")
        sz_el.set(qn("w:val"), str(int(size * 2)))
        szCs_el = OxmlElement("w:szCs")
        szCs_el.set(qn("w:val"), str(int(size * 2)))
        rPr.append(sz_el)
        rPr.append(szCs_el)
        if bold:
            for el in rPr.findall(qn("w:b")):
                rPr.remove(el)
            rPr.insert(0, OxmlElement("w:b"))

    def _para(sb=0, sa=0, align=WD_ALIGN_PARAGRAPH.LEFT):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(sb)
        p.paragraph_format.space_after = Pt(sa)
        p.paragraph_format.alignment = align
        return p

    def _r(p, text, bold=False, size=SZ):
        run = p.add_run(text)
        run.bold = bold
        run.italic = False
        run.font.name = FONT
        run.font.size = Pt(size)
        _apply_font(run._r.get_or_add_rPr(), bold=bold, size=size)
        return run

    def _link(p, text, url):
        rid = p.part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
        hl = OxmlElement("w:hyperlink")
        hl.set(qn("r:id"), rid)
        run_el = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        _apply_font(rPr)
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rPr.append(u)
        col = OxmlElement("w:color")
        col.set(qn("w:val"), LINK_COLOR)
        rPr.append(col)
        run_el.append(rPr)
        t = OxmlElement("w:t")
        t.text = text
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        run_el.append(t)
        hl.append(run_el)
        p._p.append(hl)

    # ── Name / contact header ────────────────────────────────────────
    p0 = _para(sb=0, sa=1, align=WD_ALIGN_PARAGRAPH.CENTER)
    run = p0.add_run(signer["name"].upper())
    _apply_font(run._r.get_or_add_rPr(), bold=True, size=NAME_SZ)
    run.bold = True
    run.italic = False
    run.font.size = Pt(NAME_SZ)
    run.font.name = FONT

    p0 = _para(sb=0, sa=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    contact_bits = [b for b in [signer.get("phone"), signer.get("email")] if b]
    _r(p0, " | ".join(contact_bits))
    for k in ("linkedin", "github", "website"):
        if signer.get(k) and signer[k].get("url"):
            _r(p0, " | ")
            _link(p0, signer[k]["text"], signer[k]["url"])

    # ── Date / greeting ──────────────────────────────────────────────
    if data.get("date"):
        p0 = _para(sb=0, sa=8)
        _r(p0, data["date"])
    if data.get("greeting"):
        p0 = _para(sb=0, sa=8)
        _r(p0, data["greeting"])

    # ── Body paragraphs ──────────────────────────────────────────────
    for i, paragraph in enumerate(body):
        sa = 8 if i < len(body) - 1 else 0
        p0 = _para(sb=0, sa=sa, align=WD_ALIGN_PARAGRAPH.LEFT)
        _r(p0, paragraph)

    # ── Closing / signature ──────────────────────────────────────────
    p0 = _para(sb=16, sa=2)
    _r(p0, data.get("closing", "Sincerely,"))
    p0 = _para(sb=0, sa=1)
    _r(p0, signer["name"], bold=True)
    if signer.get("phone"):
        _r(_para(sb=0, sa=0), signer["phone"])
    if signer.get("email"):
        _r(_para(sb=0, sa=0), signer["email"])

    return doc


def main():
    ap = argparse.ArgumentParser(description="Generate a cover letter DOCX from a data file.")
    ap.add_argument("--data", required=True, help="Path to the cover letter data JSON.")
    ap.add_argument("--config", default=None, help="Path to a style config (YAML or JSON).")
    ap.add_argument("--out-dir", default=None, help="Override output directory.")
    args = ap.parse_args()

    with open(os.path.expanduser(args.data), "r") as f:
        data = json.load(f)
    cfg = load_config(args.config)

    if "signer" not in data or "name" not in data.get("signer", {}):
        sys.exit("data file must contain signer.name")

    doc = build_cover_letter(data, cfg)

    company = data.get("company", "Company")
    company_slug = _safe_component(company)
    name_slug = _slug_name(data["signer"]["name"])
    fname = cfg["cover_letter_filename_pattern"].format(name=name_slug, company=company_slug)
    if not fname.endswith(".docx"):
        fname += ".docx"
    out_dir = args.out_dir or cfg["output_dir"].format(company=company_slug)
    out_dir = os.path.expanduser(out_dir)
    os.makedirs(out_dir, exist_ok=True)
    out = os.path.join(out_dir, fname)
    doc.save(out)
    print(f"Saved: {out}")


if __name__ == "__main__":
    main()
